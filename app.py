# ==============================
# Production-Ready RAG System
# Improved version of your code
# ==============================

import streamlit as st
import openai
import numpy as np
import faiss
import os
import pickle
from pathlib import Path
from typing import List, Dict, Tuple

# Document processing
from pypdf import PdfReader
from docx import Document
import pandas as pd
from bs4 import BeautifulSoup

# ==============================
# Configuration
# ==============================

EMBEDDING_MODEL = "text-embedding-3-small"
CHAT_MODEL = "gpt-4o-mini"
CHUNK_SIZE = 1200  # characters
CHUNK_OVERLAP = 200
INDEX_PATH = "faiss_index.bin"
META_PATH = "chunks_metadata.pkl"

SUPPORTED_TYPES = ["txt", "md", "pdf", "docx", "csv", "xlsx", "html"]


# ==============================
# Utils
# ==============================


def get_api_key() -> str | None:
    if "OPENAI_API_KEY" in st.secrets:
        return st.secrets["OPENAI_API_KEY"]

    if os.getenv("OPENAI_API_KEY"):
        return os.getenv("OPENAI_API_KEY")

    return None


# ==============================
# Document Readers
# ==============================


def read_txt(file) -> str:
    return file.read().decode("utf-8", errors="ignore")



def read_pdf(file) -> str:
    reader = PdfReader(file)
    parts = []
    for i, page in enumerate(reader.pages):
        txt = page.extract_text()
        if txt:
            parts.append(f"[Page {i+1}]\n{txt}")
    return "\n\n".join(parts)



def read_docx(file) -> str:
    doc = Document(file)
    paragraphs = []

    for p in doc.paragraphs:
        if p.text.strip():
            paragraphs.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)



def read_csv(file) -> str:
    df = pd.read_csv(file)
    return df.to_markdown(index=False)



def read_excel(file) -> str:
    xlsx = pd.ExcelFile(file)
    parts = []

    for sheet in xlsx.sheet_names:
        df = pd.read_excel(xlsx, sheet_name=sheet)
        parts.append(f"### Sheet: {sheet}")
        parts.append(df.to_markdown(index=False))

    return "\n\n".join(parts)



def read_html(file) -> str:
    soup = BeautifulSoup(file.read(), "html.parser")

    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    text = soup.get_text(separator="\n")
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    return "\n".join(lines)



def read_file(file) -> Tuple[str, str]:
    name = file.name.lower()

    if name.endswith(".pdf"):
        return read_pdf(file), "PDF"
    elif name.endswith(".docx"):
        return read_docx(file), "Word"
    elif name.endswith(".csv"):
        return read_csv(file), "CSV"
    elif name.endswith(".xlsx"):
        return read_excel(file), "Excel"
    elif name.endswith(".html") or name.endswith(".htm"):
        return read_html(file), "HTML"
    else:
        return read_txt(file), "Text"


# ==============================
# Chunking (IMPROVED)
# ==============================


def chunk_text_with_metadata(text: str, source: str) -> List[Dict]:
    """Character-based chunking with metadata."""

    chunks = []
    start = 0

    while start < len(text):
        end = start + CHUNK_SIZE
        chunk_text = text[start:end]

        if chunk_text.strip():
            chunks.append(
                {
                    "text": chunk_text,
                    "source": source,
                }
            )

        start += CHUNK_SIZE - CHUNK_OVERLAP

    return chunks


# ==============================
# Embeddings (CACHED + NORMALIZED)
# ==============================


@st.cache_data(show_spinner=False)
def get_embeddings_cached(texts: Tuple[str, ...], api_key: str) -> np.ndarray:
    client = openai.OpenAI(api_key=api_key)

    all_embeddings = []
    batch_size = 100

    for i in range(0, len(texts), batch_size):
        batch = list(texts[i : i + batch_size])
        resp = client.embeddings.create(model=EMBEDDING_MODEL, input=batch)
        all_embeddings.extend([d.embedding for d in resp.data])

    arr = np.array(all_embeddings).astype("float32")

    # IMPORTANT: cosine normalization
    faiss.normalize_L2(arr)
    return arr


# ==============================
# FAISS Index
# ==============================


def create_faiss_index(embeddings: np.ndarray) -> faiss.Index:
    dim = embeddings.shape[1]
    index = faiss.IndexFlatIP(dim)  # cosine via normalized vectors
    index.add(embeddings)
    return index



def save_index(index: faiss.Index, metadata: List[Dict]):
    faiss.write_index(index, INDEX_PATH)
    with open(META_PATH, "wb") as f:
        pickle.dump(metadata, f)



def load_index() -> Tuple[faiss.Index | None, List[Dict]]:
    if Path(INDEX_PATH).exists() and Path(META_PATH).exists():
        index = faiss.read_index(INDEX_PATH)
        with open(META_PATH, "rb") as f:
            metadata = pickle.load(f)
        return index, metadata
    return None, []


# ==============================
# Retrieval
# ==============================


def search_similar(
    query: str,
    index: faiss.Index,
    metadata: List[Dict],
    api_key: str,
    top_k: int = 3,
) -> List[Dict]:

    query_emb = get_embeddings_cached((query,), api_key)
    faiss.normalize_L2(query_emb)

    distances, indices = index.search(query_emb, top_k)

    results = []
    for idx in indices[0]:
        if idx < len(metadata):
            results.append(metadata[idx])

    return results


# ==============================
# Generation (ANTI-HALLUCINATION)
# ==============================


def generate_answer(query: str, contexts: List[Dict], api_key: str) -> str:
    client = openai.OpenAI(api_key=api_key)

    context_text = "\n\n---\n\n".join(c["text"] for c in contexts)

    system_prompt = """You are a strict RAG assistant.

Use ONLY the provided context.
If the answer is not explicitly in the context, say:
"I couldn't find relevant information in the documents."

Always be factual and concise.
"""

    user_prompt = f"""Context:
{context_text}

Question: {query}
Answer:"""

    resp = client.chat.completions.create(
        model=CHAT_MODEL,
        temperature=0.2,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    )

    return resp.choices[0].message.content

# ==============================
# Production-Ready RAG System (Improved UI + Core)
# ==============================

import streamlit as st
import openai
import numpy as np
import faiss
import os
import pickle
from pathlib import Path
from typing import List, Dict, Tuple

# Document processing
from pypdf import PdfReader
from docx import Document
import pandas as pd
from bs4 import BeautifulSoup

# ==============================
# Configuration
# ==============================

EMBEDDING_MODEL = "text-embedding-3-small"
CHAT_MODEL = "gpt-4o-mini"
CHUNK_SIZE = 1200  # characters (more robust than words)
CHUNK_OVERLAP = 200
INDEX_PATH = "faiss_index.bin"
META_PATH = "chunks_metadata.pkl"

SUPPORTED_TYPES = ["txt", "md", "pdf", "docx", "csv", "xlsx", "html"]


# ==============================
# Utils
# ==============================


def get_api_key() -> str | None:
    if "OPENAI_API_KEY" in st.secrets:
        return st.secrets["OPENAI_API_KEY"]
    if os.getenv("OPENAI_API_KEY"):
        return os.getenv("OPENAI_API_KEY")
    return None


# ==============================
# Document Readers
# ==============================


def read_txt(file) -> str:
    return file.read().decode("utf-8", errors="ignore")



def read_pdf(file) -> str:
    reader = PdfReader(file)
    parts = []
    for i, page in enumerate(reader.pages):
        txt = page.extract_text()
        if txt:
            parts.append(f"[Page {i+1}]\n{txt}")
    return "\n\n".join(parts)



def read_docx(file) -> str:
    doc = Document(file)
    paragraphs = []

    for p in doc.paragraphs:
        if p.text.strip():
            paragraphs.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)



def read_csv(file) -> str:
    df = pd.read_csv(file)
    return df.to_markdown(index=False)



def read_excel(file) -> str:
    xlsx = pd.ExcelFile(file)
    parts = []

    for sheet in xlsx.sheet_names:
        df = pd.read_excel(xlsx, sheet_name=sheet)
        parts.append(f"### Sheet: {sheet}")
        parts.append(df.to_markdown(index=False))

    return "\n\n".join(parts)



def read_html(file) -> str:
    soup = BeautifulSoup(file.read(), "html.parser")

    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    text = soup.get_text(separator="\n")
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    return "\n".join(lines)



def read_file(file) -> Tuple[str, str]:
    name = file.name.lower()

    if name.endswith(".pdf"):
        return read_pdf(file), "PDF"
    elif name.endswith(".docx"):
        return read_docx(file), "Word"
    elif name.endswith(".csv"):
        return read_csv(file), "CSV"
    elif name.endswith(".xlsx"):
        return read_excel(file), "Excel"
    elif name.endswith(".html") or name.endswith(".htm"):
        return read_html(file), "HTML"
    else:
        return read_txt(file), "Text"


# ==============================
# Chunking with metadata
# ==============================


def chunk_text_with_metadata(text: str, source: str) -> List[Dict]:
    chunks = []
    start = 0

    while start < len(text):
        end = start + CHUNK_SIZE
        chunk_txt = text[start:end]

        if chunk_txt.strip():
            chunks.append({"text": chunk_txt, "source": source})

        start += CHUNK_SIZE - CHUNK_OVERLAP

    return chunks


# ==============================
# Embeddings (cached + normalized)
# ==============================


@st.cache_data(show_spinner=False)
def get_embeddings_cached(texts: Tuple[str, ...], api_key: str) -> np.ndarray:
    client = openai.OpenAI(api_key=api_key)

    all_embeddings = []
    batch_size = 100

    for i in range(0, len(texts), batch_size):
        batch = list(texts[i : i + batch_size])
        resp = client.embeddings.create(model=EMBEDDING_MODEL, input=batch)
        all_embeddings.extend([d.embedding for d in resp.data])

    arr = np.array(all_embeddings).astype("float32")
    faiss.normalize_L2(arr)
    return arr


# ==============================
# FAISS index persistence
# ==============================


def create_faiss_index(embeddings: np.ndarray) -> faiss.Index:
    dim = embeddings.shape[1]
    index = faiss.IndexFlatIP(dim)
    index.add(embeddings)
    return index



def save_index(index: faiss.Index, metadata: List[Dict]):
    faiss.write_index(index, INDEX_PATH)
    with open(META_PATH, "wb") as f:
        pickle.dump(metadata, f)



def load_index() -> Tuple[faiss.Index | None, List[Dict]]:
    if Path(INDEX_PATH).exists() and Path(META_PATH).exists():
        index = faiss.read_index(INDEX_PATH)
        with open(META_PATH, "rb") as f:
            metadata = pickle.load(f)
        return index, metadata
    return None, []


# ==============================
# Retrieval
# ==============================


def search_similar(query: str, index: faiss.Index, metadata: List[Dict], api_key: str, top_k: int = 3) -> List[Dict]:
    query_emb = get_embeddings_cached((query,), api_key)
    faiss.normalize_L2(query_emb)

    _, indices = index.search(query_emb, top_k)

    results = []
    for idx in indices[0]:
        if idx < len(metadata):
            results.append(metadata[idx])
    return results


# ==============================
# Generation (anti-hallucination)
# ==============================


def generate_answer(query: str, contexts: List[Dict], api_key: str) -> str:
    client = openai.OpenAI(api_key=api_key)

    context_text = "\n\n---\n\n".join(c["text"] for c in contexts)

    system_prompt = (
        "You are a strict RAG assistant. Use ONLY the provided context. "
        "If the answer is missing, say: I couldn't find relevant information in the documents."
    )

    user_prompt = f"Context:\n{context_text}\n\nQuestion: {query}\nAnswer:"

    resp = client.chat.completions.create(
        model=CHAT_MODEL,
        temperature=0.2,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    )

    return resp.choices[0].message.content


# ==============================
# Streamlit UI (FIXED & IMPROVED)
# ==============================


def main():
    st.set_page_config(page_title="RAG Production System", page_icon="📚", layout="wide")

    st.title("📚 RAG Document Q&A System")
    st.caption("Production-ready pipeline with persistent FAISS and metadata")

    api_key = get_api_key()
    if not api_key:
        st.error("OpenAI API key not found")
        return

    # Load index once
    if "index" not in st.session_state:
        idx, meta = load_index()
        st.session_state.index = idx
        st.session_state.metadata = meta
        st.session_state.chat_history = []

    # ================= Sidebar =================
    with st.sidebar:
        st.header("📄 Documents")

        uploaded_files = st.file_uploader(
            "Upload files",
            type=SUPPORTED_TYPES,
            accept_multiple_files=True,
        )

        top_k = st.slider("Top-K retrieval", 1, 10, 3)

        if st.button("🔄 Process Documents", type="primary"):
            if not uploaded_files:
                st.warning("Please upload files first")
                return

            all_chunks: List[Dict] = []

            with st.spinner("Reading and chunking..."):
                for file in uploaded_files:
                    content, _ = read_file(file)
                    chunks = chunk_text_with_metadata(content, file.name)
                    all_chunks.extend(chunks)

            texts = tuple(c["text"] for c in all_chunks)

            with st.spinner("Creating embeddings..."):
                embeddings = get_embeddings_cached(texts, api_key)
                index = create_faiss_index(embeddings)

            save_index(index, all_chunks)

            st.session_state.index = index
            st.session_state.metadata = all_chunks
            st.success(f"✅ Processed {len(all_chunks)} chunks")

        if st.button("🗑️ Clear index"):
            if Path(INDEX_PATH).exists():
                Path(INDEX_PATH).unlink()
            if Path(META_PATH).exists():
                Path(META_PATH).unlink()
            st.session_state.index = None
            st.session_state.metadata = []
            st.rerun()

    # ================= Chat =================

    if st.session_state.index is None:
        st.info("Upload and process documents from the sidebar.")
        return

    # Show history
    for msg in st.session_state.chat_history:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    # Chat input
    if query := st.chat_input("Ask a question about your documents..."):
        st.session_state.chat_history.append({"role": "user", "content": query})

        with st.chat_message("assistant"):
            with st.spinner("Searching..."):
                results = search_similar(
                    query,
                    st.session_state.index,
                    st.session_state.metadata,
                    api_key,
                    top_k=top_k,
                )

                answer = generate_answer(query, results, api_key)

            st.markdown(answer)

            with st.expander("📎 Sources"):
                for r in results:
                    st.markdown(f"**Source:** {r['source']}")
                    preview = r["text"][:400] + "..."
                    st.markdown(f"> {preview}")

            st.session_state.chat_history.append(
                {"role": "assistant", "content": answer}
            )


if __name__ == "__main__":
    main()
